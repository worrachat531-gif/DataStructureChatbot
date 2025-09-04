# document_reader.py

import os
from typing import Optional, List
from docx import Document


class DocumentReader:
    """Class for reading and processing Word documents"""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def validate_file(self) -> bool:
        """Validate that the file exists and is accessible"""
        if not os.path.exists(self.file_path):
            print(f"Error: File {self.file_path} not found!")
            return False

        if not self.file_path.lower().endswith('.docx'):
            print(f"Error: File {self.file_path} is not a .docx file!")
            return False

        return True

    def read_paragraphs(self, doc: Document) -> List[str]:
        """Extract text from document paragraphs"""
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        return content

    def read_tables(self, doc: Document) -> List[str]:
        """Extract text from document tables"""
        content = []
        for i, table in enumerate(doc.tables, 1):
            content.append(f"\n--- TABLE {i} ---")
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                content.append(" | ".join(row_data))
            content.append(f"--- END TABLE {i} ---\n")
        return content

    def read_document(self) -> Optional[str]:
        """Read and return the complete contents of the document"""
        if not self.validate_file():
            return None

        try:
            doc = Document(self.file_path)
            content = []

            # Read all paragraphs
            content.extend(self.read_paragraphs(doc))

            # Read all tables
            content.extend(self.read_tables(doc))

            return '\n'.join(content)

        except Exception as e:
            print(f"Error reading file: {e}")
            return None

    def get_document_summary(self) -> str:
        """
        Get a comprehensive version of the document content for API usage.
        (Previously was a summary, now reads full document content).
        """
        if not self.validate_file():
            return "Error: Could not read the dataset file."

        try:
            doc = Document(self.file_path)
            content = []

            # Read all paragraphs
            content.extend(self.read_paragraphs(doc))

            # Read all tables, including all rows
            content.extend(self.read_tables(doc))

            summary = '\n'.join(content)

            if not summary.strip():
                 return "Error: Document content is empty."

            return summary

        except Exception as e:
            print(f"Error reading file: {e}")
            return "Error: Could not read the dataset file."


def read_kmutnb_dataset(file_path: str = "C:\kmutnb_chatbot\workaw\DataSetDataStructure.docx") -> str:
    """
    Convenience function to read the KMUTNB dataset

    Args:
        file_path: Path to the Word document

    Returns:
        String content of the document
    """
    reader = DocumentReader(file_path)
    content = reader.read_document()

    if content is None:
        return "Error: Could not read the dataset file."

    return content


def get_kmutnb_summary(file_path: str = "C:\kmutnb_chatbot\workaw\DataSetDataStructure.docx") -> str:
    """
    Get a summary of the KMUTNB dataset for API usage.
    This function is now modified to return the full content, not just a summary.

    Args:
        file_path: Path to the Word document

    Returns:
        String content of the document
    """
    reader = DocumentReader(file_path)
    return reader.get_document_summary()